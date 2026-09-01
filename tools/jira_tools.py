# SPDX-License-Identifier: Apache-2.0
# ============================================================
# JIRA INTEGRATION TOOLS
# AiNxt Agentic Platform — Jira Cloud REST API v3
#
# Enterprise audit trail, ticket lifecycle, and agent linkage.
#
# Env vars:
#   JIRA_URL        — https://your-org.atlassian.net
#   JIRA_PROJECT    — default project key (e.g. AiNxt)
# ============================================================

import os
import json
import base64
import threading
import urllib.request
import urllib.parse
import urllib.error
from datetime import datetime, timezone
from typing import Optional, List

from core.logger import logger


# ============================================================
# DEFAULT PROJECT
# ============================================================

def _default_project() -> str:
    return os.environ.get("JIRA_PROJECT", "AINRPY")


# ============================================================
# AUTH HELPERS
# ============================================================

# Per-thread credential override — set by JiraAdapter after the ConnectorEngine
# hands it an already-decrypted per-user token. Mirrors gitlab_tools.set_token().
# A thread-local avoids process-wide os.environ mutation, which is unsafe across
# concurrent workers.
_thread_local = threading.local()


def set_credentials(email: str, token: str) -> None:
    """Set per-thread Atlassian credentials for subsequent Jira calls.

    Called by connectors/adapters/jira.py so the connector path reuses this
    module instead of duplicating the HTTP client. ALWAYS pair with
    clear_credentials() in a finally block — a leaked credential under a
    thread-pooled server means one user's token serving another user's request.
    """
    _thread_local.email = (email or "").strip()
    _thread_local.token = (token or "").strip()


def clear_credentials() -> None:
    """Clear per-thread Atlassian credentials. Never let them outlive a request."""
    _thread_local.email = ""
    _thread_local.token = ""


def _auth_for_user(user_id: str = "", user_email: str = "") -> tuple:
    """
    Return (auth_email, api_token) for Jira Basic-auth calls.

    Resolution order:
      0. Thread-local credentials injected by set_credentials() (connector path)
      1. Per-user token stored in DB (user_id or email lookup)
      2. Service account: JIRA_EMAIL + JIRA_API_TOKEN env vars
      3. Raises PermissionError — pipeline will fail with a clear message
    """
    tl_email = getattr(_thread_local, "email", "")
    tl_token = getattr(_thread_local, "token", "")
    if tl_email and tl_token:
        return tl_email, tl_token

    try:
        from core.platform_credentials import get_atlassian_creds
        return get_atlassian_creds(user_id=user_id, email=user_email)
    except PermissionError:
        pass

    svc_email = os.environ.get("JIRA_EMAIL", "").strip()
    svc_token = os.environ.get("JIRA_API_TOKEN", "").strip()
    if svc_email and svc_token:
        return svc_email, svc_token

    raise PermissionError(
        "No JIRA credentials available. Set JIRA_EMAIL + JIRA_API_TOKEN in .env "
        "(service account), or store a personal Atlassian token via the user profile."
    )


def _jira_base() -> str:
    return os.environ.get("JIRA_URL", "").rstrip("/")


def _make_opener():
    """Build a urllib opener for direct Jira calls, honoring enterprise forward proxies."""
    import ssl

    # TLS certificate verification is always enforced (CWE-599). When
    # REQUESTS_CA_BUNDLE / SSL_CERT_FILE names a bundle we verify against it;
    # otherwise we verify against the system trust store. There is deliberately
    # no unverified fallback — an environment without a usable trust anchor must
    # be provisioned with one rather than silently accepting forged certificates.
    _ca_bundle = os.environ.get("REQUESTS_CA_BUNDLE") or os.environ.get("SSL_CERT_FILE") or ""
    ctx = ssl.create_default_context(cafile=_ca_bundle or None)

    handlers = [urllib.request.HTTPSHandler(context=ctx)]
    forward_proxy = (
        os.environ.get("HTTPS_PROXY")
        or os.environ.get("https_proxy")
        or os.environ.get("FORWARD_PROXY_URL")
        or ""
    )
    if forward_proxy:
        handlers.insert(0, urllib.request.ProxyHandler({"https": forward_proxy, "http": forward_proxy}))

    return urllib.request.build_opener(*handlers)


def _request(method: str, path: str, payload: Optional[dict] = None,
             auth_email: str = "", auth_token: str = "") -> dict:
    """
    Send a Jira REST API call.

    In production (LLM_PROXY_URL set): routes through the LLM proxy server LLM proxy.
    Jira is Atlassian Cloud — only reachable from the LLM proxy server, not from the gateway.

    In local dev (LLM_PROXY_URL unset): calls Jira directly for convenience.
    """
    from core.circuit_breaker import get_breaker
    from core.retry import retry_llm

    proxy_url = os.environ.get("LLM_PROXY_URL", "").rstrip("/")

    if proxy_url:
        # ── Production path: relay through the LLM proxy server LLM proxy ──
        import httpx

        req_body: dict = {"service": "jira", "method": method, "path": path}
        if payload is not None:
            req_body["body"] = payload
        if auth_email:
            req_body["email"] = auth_email
        if auth_token:
            req_body["token"] = auth_token

        # ── Correlation ID propagation ─────────────────────────────────────────
        # Inject request_id / chat_id from thread-local logger context so that
        # the llm_proxy service logs every Jira API call under the same
        # identifiers as the originating gateway /ask or SDLC pipeline request.
        try:
            from core.logger import get_request_id, get_chat_id
            _rid = get_request_id()
            _cid = get_chat_id()
            if _rid and _rid != "-":
                req_body["request_id"] = _rid
            if _cid and _cid != "-":
                req_body["chat_id"] = _cid
        except Exception:
            pass
        # ──────────────────────────────────────────────────────────────────────

        def _do_proxy():
            from core.proxy_tool_use import llm_proxy_headers as _lph
            resp = httpx.post(
                f"{proxy_url}/atlassian/proxy",
                json=req_body,
                headers=_lph(),
                timeout=30.0,
            )
            if 400 <= resp.status_code < 500 and resp.status_code != 429:
                raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
            resp.raise_for_status()
            # A successful mutation (e.g. label PUT) returns 204 No Content with
            # an empty body — resp.json() would raise on that. Guard it.
            return resp.json() if resp.content else {}

        try:
            return get_breaker("jira").call(
                lambda: retry_llm(_do_proxy, max_attempts=3, base_delay=1.0)
            )
        except RuntimeError as cb_err:
            logger.warning(f"Jira {method} {path} circuit OPEN: {cb_err}")
            raise Exception(str(cb_err))

    else:
        # ── Local dev fallback: call Jira directly ──
        _local_email = auth_email or os.environ.get("JIRA_EMAIL", "")
        _local_token = auth_token or os.environ.get("JIRA_API_TOKEN", "")
        if not _local_email or not _local_token:
            raise PermissionError(
                f"No JIRA credentials for direct call to {path}. "
                "Set JIRA_EMAIL + JIRA_API_TOKEN in .env."
            )
        url  = f"{_jira_base()}{path}"
        data = json.dumps(payload).encode() if payload else None

        def _do_direct():
            req = urllib.request.Request(
                url,
                data=data,
                headers={
                    "Authorization": _auth_header(email=_local_email, token=_local_token),
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                },
                method=method,
            )
            try:
                with _make_opener().open(req, timeout=60) as r:
                    body = r.read()
                    return json.loads(body) if body else {}
            except urllib.error.HTTPError as e:
                body = e.read().decode(errors="replace")
                if 400 <= e.code < 500 and e.code != 429:
                    raise Exception(f"HTTP {e.code}: {body[:300]}")
                logger.error(f"Jira {method} {path} → HTTP {e.code}: {body[:500]}")
                raise Exception(f"HTTP {e.code}: {body[:300]}")

        try:
            return get_breaker("jira").call(
                lambda: retry_llm(_do_direct, max_attempts=3, base_delay=1.0)
            )
        except RuntimeError as cb_err:
            logger.warning(f"Jira {method} {path} circuit OPEN: {cb_err}")
            raise Exception(str(cb_err))


def _put(path: str, payload: dict) -> dict:
    return _request("PUT", path, payload)


def _adf_text(text: str) -> dict:
    """Atlassian Document Format paragraph node."""
    return {
        "type": "doc", "version": 1,
        "content": [{"type": "paragraph", "content": [{"type": "text", "text": text}]}],
    }

def _adf_doc(blocks: list) -> dict:
    """
    Build a structured Atlassian Document Format doc from a list of block
    descriptors. Supported block kinds:

        {"kind": "heading",    "text": str, "level": int}
        {"kind": "paragraph",  "text": str}
        {"kind": "bulletList", "items": [str, ...]}
        {"kind": "table",      "headers": [str, ...], "rows": [[str, ...], ...]}

    Unknown kinds are skipped defensively. ADF text nodes cannot be empty, so
    blank cell/heading/paragraph text is coerced to a single space. Always
    returns a doc with at least one content node.
    """
    def _text_node(value) -> dict:
        text = str(value) if value not in (None, "") else " "
        if not text:
            text = " "
        return {"type": "text", "text": text}

    content = []
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        kind = block.get("kind")

        if kind == "heading":
            try:
                level = int(block.get("level", 1))
            except (TypeError, ValueError):
                level = 1
            level = max(1, min(6, level))
            content.append({
                "type": "heading",
                "attrs": {"level": level},
                "content": [_text_node(block.get("text"))],
            })

        elif kind == "paragraph":
            content.append({
                "type": "paragraph",
                "content": [_text_node(block.get("text"))],
            })

        elif kind == "bulletList":
            items = block.get("items") or []
            list_content = [
                {
                    "type": "listItem",
                    "content": [{"type": "paragraph", "content": [_text_node(item)]}],
                }
                for item in items
            ]
            if list_content:
                content.append({"type": "bulletList", "content": list_content})

        elif kind == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            table_rows = []
            if headers:
                table_rows.append({
                    "type": "tableRow",
                    "content": [
                        {
                            "type": "tableHeader",
                            "attrs": {},
                            "content": [{"type": "paragraph", "content": [_text_node(h)]}],
                        }
                        for h in headers
                    ],
                })
            for row in rows:
                table_rows.append({
                    "type": "tableRow",
                    "content": [
                        {
                            "type": "tableCell",
                            "attrs": {},
                            "content": [{"type": "paragraph", "content": [_text_node(cell)]}],
                        }
                        for cell in row
                    ],
                })
            if table_rows:
                content.append({
                    "type": "table",
                    "attrs": {"isNumberColumnEnabled": False, "layout": "default"},
                    "content": table_rows,
                })
        # else: unknown kind — skip

    if not content:
        content.append({"type": "paragraph", "content": [_text_node(None)]})

    return {"type": "doc", "version": 1, "content": content}

# ============================================================
# JIRA CREATE ISSUE
# ============================================================

def jira_create_issue(
        summary: str,
        description: str,
        project: str = "",
        priority: str = "Medium",
        issue_type: str = "Bug",
        user_id: str = "",
        user_email: str = "",
        repo_name: str = "",
) -> str:
    """
    Create a Jira issue.
    Returns the issue URL as a string.

    project key resolution:
      1. Explicit ``project`` arg
      2. Product linked to ``repo_name`` (product_repos → products.jira_project_key)
      3. JIRA_PROJECT env var

    Auth resolution:
      user_id / user_email → user_tokens (atlassian); raises PermissionError if not found
    """
    try:
        # Resolve project key from product if not explicitly given
        if not project and repo_name:
            try:
                from core.platform_credentials import get_product_for_repo
                ctx = get_product_for_repo(repo_name)
                project = ctx.get("jira_project_key", "")
            except Exception:
                pass
        proj = (project or _default_project()).upper()

        try:
            from core.prompt_sanitizer import sanitize as _san
            summary     = _san(summary)
            description = _san(description)
        except Exception:
            pass

        auth_email, auth_token = _auth_for_user(user_id, user_email)
        payload = {
            "fields": {
                "project": {"key": proj},
                "summary": summary,
                "description": {
                    "type": "doc",
                    "version": 1,
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": description}],
                        }
                    ],
                },
                "issuetype": {"name": issue_type},
                "priority": {"name": priority},
            }
        }
        result = _request("POST", "/rest/api/3/issue", payload, auth_email=auth_email, auth_token=auth_token)
        issue_key = result.get("key", "UNKNOWN")
        url = f"{_jira_base()}/browse/{issue_key}"
        logger.info(f"Jira issue created → {issue_key} (project={proj})")
        return url
    except Exception as e:
        logger.error(f"jira_create_issue failed: {e}")
        return f"Error creating Jira issue: {e}"


# ============================================================
# JIRA SEARCH (JQL)
# ============================================================

_DEFAULT_SEARCH_FIELDS = "summary,status,assignee,reporter,priority,issuetype,created,updated"


def _normalize_issue(item: dict) -> dict:
    """Flatten a raw Jira issue into the compact shape returned to callers."""
    fields    = item.get("fields") or {}
    assignee  = fields.get("assignee")  or {}
    reporter  = fields.get("reporter")  or {}
    status    = fields.get("status")    or {}
    priority  = fields.get("priority")  or {}
    issuetype = fields.get("issuetype") or {}
    key = item.get("key", "")
    return {
        "id":         item.get("id", ""),
        "key":        key,
        "summary":    fields.get("summary", ""),
        "status":     status.get("name", ""),
        "assignee":   assignee.get("displayName", ""),
        "reporter":   reporter.get("displayName", ""),
        "priority":   priority.get("name", ""),
        "issue_type": issuetype.get("name", ""),
        "created":    fields.get("created", ""),
        "updated":    fields.get("updated", ""),
        "url":        f"{_jira_base()}/browse/{key}" if key else "",
    }


def jira_search_issues(jql: str, fields: str = "", limit: int = 25,
                       cursor: str = "",
                       user_id: str = "", user_email: str = "") -> dict:
    """
    Search Jira issues with JQL. Returns {"issues": [...], "next_cursor": str|None}.

    Uses POST /rest/api/3/search/jql. The old GET /rest/api/3/search was REMOVED
    from Jira Cloud, so this is the only supported search endpoint.

    Pagination is cursor-based (nextPageToken), not startAt/maxResults offsets —
    pass the returned next_cursor back as `cursor` to fetch the following page.
    Note this endpoint does NOT return a total count; use jira_count_issues().
    """
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    try:
        max_results = max(1, min(int(limit), 100))
    except (TypeError, ValueError):
        max_results = 25

    field_list = [f.strip() for f in (fields or _DEFAULT_SEARCH_FIELDS).split(",") if f.strip()]
    payload: dict = {
        "jql":        (jql or "").strip(),
        "maxResults": max_results,
        "fields":     field_list,
    }
    if cursor:
        payload["nextPageToken"] = cursor

    result = _request("POST", "/rest/api/3/search/jql", payload,
                      auth_email=auth_email, auth_token=auth_token)
    raw = result.get("issues") or []
    return {
        "issues":      [_normalize_issue(i) for i in raw],
        "next_cursor": result.get("nextPageToken"),
    }


def jira_count_issues(jql: str, user_id: str = "", user_email: str = "") -> dict:
    """
    Approximate issue count for a JQL query, without fetching the issues.

    Uses POST /rest/api/3/search/approximate-count — far cheaper than paging
    through results just to count them.
    """
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    result = _request("POST", "/rest/api/3/search/approximate-count",
                      {"jql": (jql or "").strip()},
                      auth_email=auth_email, auth_token=auth_token)
    return {"count": result.get("count", 0), "jql": (jql or "").strip()}


# ============================================================
# JIRA LIST ISSUES
# ============================================================

def jira_list_issues(project: str, status: str = "Open",
                     user_id: str = "", user_email: str = "") -> str:
    """
    List Jira issues for a project.
    Returns formatted string.

    Thin wrapper over jira_search_issues() so there is a single search code path.
    """
    try:
        jql = f'project = "{project.upper()}" AND status = "{status}" ORDER BY created DESC'
        page = jira_search_issues(jql, limit=20, user_id=user_id, user_email=user_email)
        issues = page.get("issues", [])
        if not issues:
            return f"No issues found in project {project} with status {status}."
        lines = [f"Issues in {project} ({status}):"]
        for issue in issues:
            lines.append(
                f"• [{issue.get('key')}] {issue.get('summary', '')} "
                f"(Priority: {issue.get('priority') or 'N/A'})"
            )
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"jira_list_issues failed: {e}")
        return f"Error listing Jira issues: {e}"


# ============================================================
# JIRA GET ISSUE
# ============================================================

def _adf_to_text(node) -> str:
    """
    Recursively convert an Atlassian Document Format (ADF) node to plain text.
    Handles doc, paragraph, text, heading, bulletList, listItem, codeBlock, blockquote.
    """
    if node is None:
        return ""
    if isinstance(node, str):
        return node
    if isinstance(node, list):
        return "\n".join(_adf_to_text(n) for n in node)
    node_type = node.get("type", "")
    text = node.get("text")
    if text is not None:
        return str(text)
    children = node.get("content", [])
    if node_type in ("doc", "paragraph", "listItem", "blockquote"):
        return " ".join(_adf_to_text(c) for c in children if c)
    if node_type in ("bulletList", "orderedList"):
        parts = []
        for i, c in enumerate(children, 1):
            prefix = f"{i}. " if node_type == "orderedList" else "• "
            parts.append(prefix + _adf_to_text(c))
        return "\n".join(parts)
    if node_type in ("heading", "codeBlock"):
        return _adf_to_text(children)
    # Default: recurse into content
    return " ".join(_adf_to_text(c) for c in children if c)


# ── Attachment text extraction (feeds the normalizer) ─────────────────────────
# Untrusted external input: parse IN-MEMORY only, never persist bytes/text to
# disk/DB/logs, extension-allowlist + size-cap to prevent zip-bomb / resource
# exhaustion, and fail-open (return "") so a bad/unsupported attachment never
# crashes the pipeline. Extracted text flows only into the transient normalizer
# LLM prompt (same trust path the description already takes). NEVER log the text.
_ATTACHMENT_TEXT_EXT = (".txt", ".md")
_ATTACHMENT_PARSE_EXT = {".txt", ".md", ".pdf", ".docx", ".xlsx"}
_ATTACHMENT_MAX_CHARS = 20000


def _download_attachment_bytes(url: str, auth_email: str, auth_token: str) -> bytes:
    """Best-effort authenticated GET of raw attachment bytes.

    Always uses the direct urllib opener (honoring the same TLS/proxy config as
    the direct path) with a bounded read. The LLM proxy relays JSON only, so
    binary download cannot reuse _request; if the opener is blocked in prod this
    fails-open to b"" and attachment text is simply omitted (description-only
    normalization is unchanged and correct). Never raises.
    """
    if not url:
        return b""
    try:
        req = urllib.request.Request(
            url,
            headers={"Authorization": _auth_header(email=auth_email, token=auth_token)},
        )
        with _make_opener().open(req, timeout=60) as r:
            return r.read(10 * 1024 * 1024)  # cap at ~10 MB
    except Exception:
        return b""


def _parse_attachment_text(filename: str, data: bytes) -> str:
    """Dispatch by lowercased extension to a safe in-memory parser.

    Each branch is wrapped in its own try/except → returns "" on any failure
    (missing lib, unparseable file). Output is truncated to _ATTACHMENT_MAX_CHARS.
    Never logs the text.
    """
    if not data:
        return ""
    name = (filename or "").lower()
    ext = ""
    dot = name.rfind(".")
    if dot != -1:
        ext = name[dot:]

    text = ""
    if ext in _ATTACHMENT_TEXT_EXT:
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:
            return ""
    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            return ""
    elif ext == ".docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text for c in row.cells if c.text]
                    if cells:
                        parts.append("\t".join(cells))
            text = "\n".join(parts)
        except Exception:
            return ""
    elif ext == ".xlsx":
        try:
            import openpyxl
            import io
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(cells):
                        parts.append("\t".join(cells))
            text = "\n".join(parts)
        except Exception:
            return ""
    else:
        return ""

    return text[:_ATTACHMENT_MAX_CHARS]


def jira_get_issue(issue_key: str, user_id: str = "", user_email: str = "") -> str:
    """
    Get details of a specific Jira issue.
    Returns formatted string.
    """
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        result = _request("GET", f"/rest/api/3/issue/{issue_key.upper()}",
                          auth_email=auth_email, auth_token=auth_token)
        fields = result.get("fields", {})
        summary = fields.get("summary", "")
        status = fields.get("status", {}).get("name", "")
        priority = fields.get("priority", {}).get("name", "")
        assignee = fields.get("assignee") or {}
        assignee_name = assignee.get("displayName", "Unassigned")
        raw_desc = fields.get("description") or {}
        desc_text = _adf_to_text(raw_desc) if isinstance(raw_desc, dict) else str(raw_desc or "")
        url = f"{_jira_base()}/browse/{issue_key.upper()}"
        return (
            f"Issue: {issue_key.upper()}\n"
            f"Summary: {summary}\n"
            f"Status: {status}\n"
            f"Priority: {priority}\n"
            f"Assignee: {assignee_name}\n"
            f"Description: {desc_text}\n"
            f"URL: {url}"
        )
    except Exception as e:
        logger.error(f"jira_get_issue failed: {e}")
        return f"Error fetching Jira issue {issue_key}: {e}"


def jira_get_issue_dict(issue_key: str, user_id: str = "", user_email: str = "") -> dict:
    """
    Get details of a specific Jira issue as a structured dict.
    Returns full description (no truncation), suitable for programmatic use.

    Returns:
        {
          "key":         str,
          "summary":     str,
          "description": str,   # full plain text, converted from ADF
          "status":      str,
          "priority":    str,
          "assignee":    str,
          "issue_type":  str,
          "url":         str,
        }
    """
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        result = _request("GET", f"/rest/api/3/issue/{issue_key.upper()}",
                          auth_email=auth_email, auth_token=auth_token)
        fields = result.get("fields", {})
        raw_desc = fields.get("description") or {}
        desc_text = _adf_to_text(raw_desc) if isinstance(raw_desc, dict) else str(raw_desc or "")
        return {
            "key":        issue_key.upper(),
            "summary":    fields.get("summary", ""),
            "description": desc_text,
            "status":     (fields.get("status") or {}).get("name", ""),
            "priority":   (fields.get("priority") or {}).get("name", "Medium"),
            "assignee":   ((fields.get("assignee") or {}) or {}).get("displayName", ""),
            "issue_type": ((fields.get("issuetype") or {}) or {}).get("name", "Story"),
            "url":        f"{_jira_base()}/browse/{issue_key.upper()}",
        }
    except Exception as e:
        logger.error(f"jira_get_issue_dict failed: {e}")
        return {}

def jira_get_issue_full(issue_key: str, user_id: str = "", user_email: str = "") -> dict:
    """
    Get full details of a Jira issue for the TICKET_NORMALIZATION stage.
    Returns a richer dict than jira_get_issue_dict: includes comments (last 10),
    attachments (names + descriptions), acceptance_criteria (custom field if present),
    labels, and epic_summary. Gives the NormalizationAgent maximum raw material.

    Returns dict with keys:
        key, summary, description, status, priority, assignee, issue_type, url,
        comments (list of {author, body, created}),
        attachments (list of {filename, content}),
        acceptance_criteria (str — custom field or ""),
        labels (list of str),
        epic_summary (str or ""),
        raw_fields (dict — full fields for any extra custom fields)
    """
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        result = _request(
            "GET",
            f"/rest/api/3/issue/{issue_key.upper()}?fields=summary,description,status,priority,"
            "assignee,issuetype,comment,attachment,labels,customfield_10014,"
            "customfield_10016,customfield_10020,customfield_10100,customfield_10101,"
            "customfield_10103,parent,&expand=renderedFields",
            auth_email=auth_email,
            auth_token=auth_token,
        )
        fields = result.get("fields", {})
        raw_desc = fields.get("description") or {}
        desc_text = _adf_to_text(raw_desc) if isinstance(raw_desc, dict) else str(raw_desc or "")

        # Comments: last 10 (most recent first from Jira's API)
        comment_data = fields.get("comment") or {}
        raw_comments = comment_data.get("comments") or []
        comments = []
        for c in raw_comments[-10:]:
            body_raw = c.get("body") or {}
            body_text = _adf_to_text(body_raw) if isinstance(body_raw, dict) else str(body_raw or "")
            author = ((c.get("author") or {}).get("displayName") or "")
            created = c.get("created", "")
            if body_text.strip():
                comments.append({"author": author, "body": body_text.strip(), "created": created})

        # Attachments
        raw_attachments = fields.get("attachment") or []
        attachments = [
            {"filename": a.get("filename", ""), "content": a.get("content", "")}
            for a in raw_attachments
        ]

        # Download + parse allowlisted attachments IN-MEMORY (never persist bytes/
        # text; never log text). Fail-open per attachment. Feeds normalization.
        n_parsed = 0
        for a in attachments:
            fn = (a.get("filename") or "").lower()
            dot = fn.rfind(".")
            ext = fn[dot:] if dot != -1 else ""
            if ext not in _ATTACHMENT_PARSE_EXT:
                continue
            raw_bytes = _download_attachment_bytes(a.get("content", ""), auth_email, auth_token)
            parsed = _parse_attachment_text(a.get("filename", ""), raw_bytes)
            if parsed:
                a["text"] = parsed
                n_parsed += 1
        attachments_text = "\n\n".join(
            f"=== {a['filename']} ===\n{a['text']}"
            for a in attachments if a.get("text")
        )

        # Acceptance criteria: try common custom field names
        ac = ""
        for cf_key in ("customfield_10016", "customfield_10020", "customfield_10100",
                       "customfield_10101", "customfield_10103"):
            cf_val = fields.get(cf_key)
            if cf_val:
                ac = _adf_to_text(cf_val) if isinstance(cf_val, dict) else str(cf_val)
                if ac.strip():
                    break

        # Labels
        labels = [str(lb) for lb in (fields.get("labels") or [])]

        # Epic summary (parent or customfield_10014)
        epic_summary = ""
        parent = fields.get("parent") or {}
        if parent:
            epic_summary = (parent.get("fields") or {}).get("summary", "") or ""
        if not epic_summary:
            cf_epic = fields.get("customfield_10014")
            if cf_epic and isinstance(cf_epic, dict):
                epic_summary = cf_epic.get("summary", "") or ""

        n_fields = len(fields)
        n_comments = len(comments)
        n_attachments = len(attachments)
        ac_present = "present" if ac.strip() else "absent"
        logger.info(
            f"[jira {issue_key.upper()}] full fetch: "
            f"fields={n_fields} comments={n_comments} attachments={n_attachments} "
            f"parsed_attachments={n_parsed} acceptance_criteria={ac_present}"
        )

        return {
            "key": issue_key.upper(),
            "summary": fields.get("summary", ""),
            "description": desc_text,
            "status": (fields.get("status") or {}).get("name", ""),
            "priority": (fields.get("priority") or {}).get("name", "Medium"),
            "assignee": ((fields.get("assignee") or {}) or {}).get("displayName", ""),
            "issue_type": ((fields.get("issuetype") or {}) or {}).get("name", "Story"),
            "url": f"{_jira_base()}/browse/{issue_key.upper()}",
            "comments": comments,
            "attachments": attachments,
            "attachments_text": attachments_text,
            "acceptance_criteria": ac.strip(),
            "labels": labels,
            "epic_summary": epic_summary,
            "raw_fields": fields,
        }
    except Exception as e:
        logger.error(f"jira_get_issue_full failed for {issue_key}: {e}")
        return {}


# ============================================================
# JIRA PROJECT / USER / METADATA LOOKUPS
# ============================================================

def jira_list_projects(limit: int = 150, search: str = "",
                       user_id: str = "", user_email: str = "") -> List[dict]:
    """
    List Jira projects visible to the authenticated user.

    Use this to resolve a project KEY before calling any project-scoped tool —
    never guess a key from an informal project name.
    """
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    try:
        max_results = max(1, min(int(limit), 150))
    except (TypeError, ValueError):
        max_results = 150

    path = f"/rest/api/3/project/search?maxResults={max_results}"
    if search:
        path += f"&query={urllib.parse.quote(str(search))}"

    result = _request("GET", path, auth_email=auth_email, auth_token=auth_token)
    return [
        {
            "key":  p.get("key", ""),
            "name": p.get("name", ""),
            "id":   p.get("id", ""),
            "url":  f"{_jira_base()}/browse/{p.get('key', '')}",
        }
        for p in (result.get("values") or [])
    ]


def jira_get_project(project_key: str, user_id: str = "", user_email: str = "") -> dict:
    """Get metadata for one Jira project — name, description, and lead."""
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    proj = (project_key or "").upper()
    result = _request("GET", f"/rest/api/3/project/{proj}",
                      auth_email=auth_email, auth_token=auth_token)
    return {
        "key":         result.get("key", ""),
        "name":        result.get("name", ""),
        "description": (result.get("description") or "")[:300],
        "lead":        (result.get("lead") or {}).get("displayName", ""),
        "url":         f"{_jira_base()}/browse/{result.get('key', '')}",
    }


def jira_get_current_user(user_id: str = "", user_email: str = "") -> dict:
    """
    Get the profile of the authenticated Jira user — accountId, name, email.

    Takes no arguments and changes nothing, so this is also the connection-test
    probe for the Jira connector (see connectors/probe.py). The returned
    account_id is what jira_assign_issue needs for self-assignment.
    """
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    result = _request("GET", "/rest/api/3/myself",
                      auth_email=auth_email, auth_token=auth_token)
    return {
        "account_id":   result.get("accountId", ""),
        "display_name": result.get("displayName", ""),
        "email":        result.get("emailAddress", ""),
        "active":       result.get("active", True),
    }


def jira_get_transitions(issue_key: str, user_id: str = "",
                         user_email: str = "") -> List[dict]:
    """
    List the status transitions currently available on an issue.

    Jira only permits transitions allowed by the project workflow from the
    issue's CURRENT status, so call this before jira_transition_issue rather
    than guessing a target status name.
    """
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    return _fetch_transitions(issue_key, auth_email, auth_token)


def _fetch_transitions(issue_key: str, auth_email: str, auth_token: str) -> List[dict]:
    """Raw transition fetch, shared by jira_get_transitions and jira_update_issue."""
    key = (issue_key or "").upper()
    result = _request("GET", f"/rest/api/3/issue/{key}/transitions",
                      auth_email=auth_email, auth_token=auth_token)
    return [
        {
            "id":        t.get("id", ""),
            "name":      t.get("name", ""),
            "to_status": ((t.get("to") or {}).get("name", "")),
        }
        for t in (result.get("transitions") or [])
    ]


def jira_list_comments(issue_key: str, limit: int = 150,
                       user_id: str = "", user_email: str = "") -> List[dict]:
    """List comments on an issue, with ADF bodies flattened to plain text."""
    auth_email, auth_token = _auth_for_user(user_id, user_email)
    key = (issue_key or "").upper()
    try:
        max_results = max(1, min(int(limit), 150))
    except (TypeError, ValueError):
        max_results = 150

    result = _request("GET", f"/rest/api/3/issue/{key}/comment?maxResults={max_results}",
                      auth_email=auth_email, auth_token=auth_token)
    comments = []
    for c in (result.get("comments") or []):
        body_raw = c.get("body") or {}
        comments.append({
            "id":      c.get("id", ""),
            "author":  (c.get("author") or {}).get("displayName", ""),
            "body":    _adf_to_text(body_raw) if isinstance(body_raw, dict) else str(body_raw or ""),
            "created": c.get("created", ""),
            "updated": c.get("updated", ""),
        })
    return comments


# ============================================================
# JIRA UPDATE ISSUE
# ============================================================

def jira_update_issue(
        issue_key: str,
        status: Optional[str] = None,
        comment: Optional[str] = None,
        assignee_account_id: Optional[str] = None,
        priority: Optional[str] = None,
        user_id: str = "",
        user_email: str = "",
) -> str:
    """
    Update an existing Jira issue — transition status, add a comment,
    change assignee, or update priority.
    Returns a confirmation string.
    """
    key = issue_key.upper()
    results = []

    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)

        # ── Transition status ──────────────────────────────
        if status:
            transitions = _fetch_transitions(key, auth_email, auth_token)
            target = next(
                (t for t in transitions
                 if t["name"].lower() == status.lower()
                 or t["to_status"].lower() == status.lower()),
                None,
            )
            if target:
                _request("POST", f"/rest/api/3/issue/{key}/transitions",
                         {"transition": {"id": target["id"]}},
                         auth_email=auth_email, auth_token=auth_token)
                results.append(f"Status → {status}")
            else:
                # List what IS available so the caller can retry with a valid name
                # instead of guessing again.
                available = ", ".join(t["name"] for t in transitions) or "none"
                results.append(f"Status '{status}' not available (valid: {available})")

        # ── Priority ───────────────────────────────────────
        if priority:
            _request("PUT", f"/rest/api/3/issue/{key}",
                     {"fields": {"priority": {"name": priority}}},
                     auth_email=auth_email, auth_token=auth_token)
            results.append(f"Priority → {priority}")

        # ── Assignee ───────────────────────────────────────
        if assignee_account_id:
            _request("PUT", f"/rest/api/3/issue/{key}",
                     {"fields": {"assignee": {"accountId": assignee_account_id}}},
                     auth_email=auth_email, auth_token=auth_token)
            results.append(f"Assignee → {assignee_account_id}")

        # ── Comment ────────────────────────────────────────
        if comment:
            _request("POST", f"/rest/api/3/issue/{key}/comment",
                     {"body": _adf_text(comment)},
                     auth_email=auth_email, auth_token=auth_token)
            results.append("Comment added")

        logger.info(f"Jira updated {key}: {', '.join(results)}")
        return f"[{key}] Updated: {', '.join(results)}"

    except Exception as e:
        logger.error(f"jira_update_issue failed: {e}")
        return f"Error updating {key}: {e}"


# ============================================================
# JIRA ADD COMMENT
# ============================================================

def jira_add_comment(issue_key: str, comment: str,
                     user_id: str = "", user_email: str = "") -> str:
    """Post a comment to a Jira issue. Returns confirmation."""
    key = issue_key.upper()
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        _request("POST", f"/rest/api/3/issue/{key}/comment",
                 {"body": _adf_text(comment)},
                 auth_email=auth_email, auth_token=auth_token)
        logger.info(f"Jira comment added → {key}")
        return f"Comment added to {key}"
    except Exception as e:
        logger.error(f"jira_add_comment failed: {e}")
        return f"Error adding comment to {key}: {e}"


# ============================================================
# JIRA LINK ISSUES
# ============================================================

def jira_link_issues(
        inward_key: str,
        outward_key: str,
        link_type: str = "Relates",
        user_id: str = "",
        user_email: str = "",
) -> str:
    """
    Link two Jira issues together. `link_type` is the Jira issue-link *type
    name* (e.g. 'Relates', 'Blocks', 'Duplicate') — not the inward/outward
    phrasing like 'relates to', which Jira rejects with a 404.
    Returns confirmation string.
    """
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        _request("POST", "/rest/api/3/issueLink", {
            "type":         {"name": link_type},
            "inwardIssue":  {"key": inward_key.upper()},
            "outwardIssue": {"key": outward_key.upper()},
        }, auth_email=auth_email, auth_token=auth_token)
        logger.info(f"Jira linked {inward_key} → {outward_key} ({link_type})")
        return f"Linked {inward_key.upper()} {link_type} {outward_key.upper()}"
    except Exception as e:
        logger.error(f"jira_link_issues failed: {e}")
        return f"Error linking issues: {e}"

# ============================================================
# GOVERNANCE EVIDENCE HELPERS
# (ADF comment writer, label, and attachment upload — used to log
#  governance approval evidence to a linked Jira Change ticket)
# ============================================================

def jira_add_comment_adf(issue_key: str, adf_doc: dict,
                         user_id: str = "", user_email: str = "") -> str:
    """
    Post a pre-built Atlassian Document Format doc (see ``_adf_doc``) as a
    comment on a Jira issue — for structured evidence (headings/tables/lists)
    that a plain-text comment can't render. Returns a confirmation string.
    """
    key = issue_key.upper()
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        _request("POST", f"/rest/api/3/issue/{key}/comment",
                 {"body": adf_doc},
                 auth_email=auth_email, auth_token=auth_token)
        logger.info(f"Jira ADF comment added → {key}")
        return f"ADF comment added to {key}"
    except Exception as e:
        logger.error(f"jira_add_comment_adf failed: {e}")
        return f"Error adding ADF comment to {key}: {e}"


def jira_add_label(issue_key: str, label: str,
                   user_id: str = "", user_email: str = "") -> str:
    """
    Add a label to a Jira issue (additive — existing labels are preserved).
    Returns a confirmation string.
    """
    key = issue_key.upper()
    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        _request("PUT", f"/rest/api/3/issue/{key}",
                 {"update": {"labels": [{"add": label}]}},
                 auth_email=auth_email, auth_token=auth_token)
        logger.info(f"Jira label added → {key}: {label}")
        return f"Label '{label}' added to {key}"
    except Exception as e:
        logger.error(f"jira_add_label failed: {e}")
        return f"Error adding label to {key}: {e}"


def jira_add_attachment(issue_key: str, filename: str, content_bytes: bytes,
                        content_type: str = "application/octet-stream",
                        user_id: str = "", user_email: str = "") -> str:
    """
    Upload a binary attachment to a Jira issue.

    ``_request()`` only carries JSON bodies, so this mirrors its prod/dev
    branch structure directly rather than reusing it:

      - production (LLM_PROXY_URL set): Jira Cloud is only reachable from
        web02, so the file is base64-encoded into a JSON body and POSTed to
        ``{LLM_PROXY_URL}/atlassian/attachment``.
      - local dev (LLM_PROXY_URL unset): a real multipart/form-data POST
        straight to Jira's attachments endpoint.

    Never raises — returns a confirmation or error string.
    """
    from core.circuit_breaker import get_breaker
    from core.retry import retry_llm

    key = issue_key.upper()
    proxy_url = os.environ.get("LLM_PROXY_URL", "").rstrip("/")

    try:
        auth_email, auth_token = _auth_for_user(user_id, user_email)
        import httpx

        if proxy_url:
            # ── Production path: relay through web02 LLM proxy ──
            from core.proxy_tool_use import llm_proxy_headers

            req_body: dict = {
                "service": "jira",
                "issue_key": key,
                "filename": filename,
                "content_b64": base64.b64encode(content_bytes).decode(),
                "content_type": content_type,
                "email": auth_email,
                "token": auth_token,
            }

            # ── Correlation ID propagation (mirrors _request) ──
            try:
                from core.logger import get_request_id, get_chat_id
                _rid = get_request_id()
                _cid = get_chat_id()
                if _rid and _rid != "-":
                    req_body["request_id"] = _rid
                if _cid and _cid != "-":
                    req_body["chat_id"] = _cid
            except Exception:
                pass

            def _do_proxy():
                resp = httpx.post(
                    f"{proxy_url}/atlassian/attachment",
                    json=req_body,
                    headers=llm_proxy_headers(),
                    timeout=30.0,
                )
                if 400 <= resp.status_code < 500 and resp.status_code != 429:
                    raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
                resp.raise_for_status()
                return resp.json()

            get_breaker("jira").call(
                lambda: retry_llm(_do_proxy, max_attempts=3, base_delay=1.0)
            )

        else:
            # ── Local dev fallback: multipart POST directly to Jira ──
            url = f"{_jira_base()}/rest/api/3/issue/{key}/attachments"

            def _do_direct():
                resp = httpx.post(
                    url,
                    headers={
                        "Authorization": _auth_header(email=auth_email, token=auth_token),
                        "X-Atlassian-Token": "no-check",
                    },
                    files={"file": (filename, content_bytes, content_type)},
                    timeout=30.0,
                )
                if 400 <= resp.status_code < 500 and resp.status_code != 429:
                    raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
                resp.raise_for_status()
                return resp.json()

            get_breaker("jira").call(
                lambda: retry_llm(_do_direct, max_attempts=3, base_delay=1.0)
            )

        logger.info("[JIRA] attachment uploaded", issue_key=key, filename=filename,
                    path=("proxy" if proxy_url else "dev"))
        return f"Attachment {filename} uploaded to {key}"

    except Exception as e:
        logger.error("[JIRA] attachment upload failed", issue_key=key, filename=filename, error=str(e))
        return f"Error uploading attachment {filename} to {key}: {e}"

# ============================================================
# JIRA LOG AGENT ACTION (ENTERPRISE AUDIT TRAIL)
# ============================================================

_AUDIT_EVENT_TYPES = {
    "agent_created":        ("Story",   "Medium", "[AUDIT] Agent Created"),
    "workflow_created":     ("Story",   "Medium", "[AUDIT] Workflow Created"),
    "workflow_executed":    ("Task",    "Low",    "[AUDIT] Workflow Executed"),
    "incident_detected":    ("Bug",     "High",   "[INCIDENT] Detected"),
    "incident_resolved":    ("Bug",     "Medium", "[INCIDENT] Resolved"),
    "code_change_proposed": ("Task",    "Medium", "[CODE] Change Proposed"),
    "code_change_merged":   ("Task",    "Low",    "[CODE] Change Merged"),
    "security_alert":       ("Bug",     "Critical", "[SECURITY] Alert"),
    "compliance_violation": ("Bug",     "High",   "[COMPLIANCE] Violation"),
    "model_cost_exceeded":  ("Task",    "High",   "[COST] Budget Exceeded"),
}

def jira_log_agent_action(
        event_type: str,
        summary: str,
        details: str,
        agent_name: str = "",
        request_id: str = "",
        project: Optional[str] = None,
        user_id: str = "",
        user_email: str = "",
) -> str:
    """
    Create an audit-trail Jira ticket for any platform event.

    event_type: one of agent_created, workflow_created, workflow_executed,
                incident_detected, incident_resolved, code_change_proposed,
                code_change_merged, security_alert, compliance_violation,
                model_cost_exceeded

    Returns the Jira issue URL or error string.
    """
    proj = project or _default_project()
    issue_type, priority, prefix = _AUDIT_EVENT_TYPES.get(
        event_type, ("Task", "Medium", f"[{event_type.upper()}]")
    )
    ts  = datetime.now(timezone.utc).isoformat()
    full_summary = f"{prefix}: {summary}"
    body = (
        f"*Event Type:* {event_type}\n"
        f"*Agent:* {agent_name or 'N/A'}\n"
        f"*Request ID:* {request_id or 'N/A'}\n"
        f"*Timestamp:* {ts}\n\n"
        f"*Details:*\n{details}"
    )
    try:
        url = jira_create_issue(
            project=proj,
            summary=full_summary,
            description=body,
            priority=priority,
            issue_type=issue_type,
            user_id=user_id,
            user_email=user_email,
        )
        logger.info(f"Jira audit logged [{event_type}] → {url}")
        return url
    except Exception as e:
        logger.error(f"jira_log_agent_action failed: {e}")
        return f"Error logging audit event: {e}"


# ============================================================
# JIRA TRANSITION / ASSIGN
# ============================================================

def jira_transition_issue(issue_key: str, status: str,
                          user_id: str = "", user_email: str = "") -> str:
    """
    Transition a Jira issue to a new status (e.g., 'In Progress', 'Done').

    Only transitions permitted by the workflow from the issue's current status
    will succeed — call jira_get_transitions() first if unsure.
    """
    return jira_update_issue(issue_key, status=status,
                             user_id=user_id, user_email=user_email)


def jira_assign_issue(issue_key: str, account_id: str,
                      user_id: str = "", user_email: str = "") -> str:
    """
    Assign a Jira issue to a user, identified by Atlassian accountId.

    Jira requires an accountId here — a display name or email will NOT work.
    Get one from jira_get_current_user() (for self-assignment).
    """
    return jira_update_issue(issue_key, assignee_account_id=account_id,
                             user_id=user_id, user_email=user_email)
