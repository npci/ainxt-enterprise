# SPDX-License-Identifier: MIT
"""
AiNxt Agentic Platform — document_tools MCP tools.

Plain-text extraction over PDF / MD / TXT / EML / CSV files under a
configured document root. Tools extract raw content; summarisation and
reasoning are the calling agent's job.

Functions exposed:
  list_documents      — list readable documents under data_dir
  extract_text        — pull plain text out of a document
  search_in_document  — find query occurrences with surrounding context

Companion server: mcp/servers/document_tools_server.py
Registered in:   mcp/registry.py:_register_tools()

Configuration (env vars):
  DOCUMENT_TOOLS_DATA_DIR  — root directory for documents (default ./data/documents)
  DOCUMENT_TOOLS_MAX_PAGES — max PDF pages to extract per call (default 50)
"""

import os
from typing import List

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None  # type: ignore[assignment]


# ── Configuration ────────────────────────────────────────────────────────────

_DATA_DIR  = os.getenv("DOCUMENT_TOOLS_DATA_DIR", "./data/documents")
_MAX_PAGES = int(os.getenv("DOCUMENT_TOOLS_MAX_PAGES", "50"))

# Extensions we can extract text from. Extended (Fix #6/#15) to cover binary Office
# documents (.docx/.xlsx/.xls) and HTML — previously these fell through to a raw
# UTF-8 read that produced garbage or failed outright.
_TEXT_EXT   = (".md", ".txt", ".csv", ".eml", ".json", ".log")
_HTML_EXT   = (".html", ".htm")
_DOCX_EXT   = (".docx",)
_XLSX_EXT   = (".xlsx",)
_XLS_EXT    = (".xls",)
# .doc/.ppt/.pptx/.rtf handled via the shared core.document_parser (G16/G17).
_CORE_EXT   = (".doc", ".ppt", ".pptx", ".rtf")
_SUPPORTED  = (".pdf",) + _TEXT_EXT + _HTML_EXT + _DOCX_EXT + _XLSX_EXT + _XLS_EXT + _CORE_EXT


# ── Format-specific extractors ────────────────────────────────────────────────

def _extract_docx(full: str) -> str:
    """Extract text from a binary Word .docx (paragraphs + tables). Fix #15."""
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx is required to read .docx files") from e
    doc = docx.Document(full)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(full: str) -> str:
    """Extract text from a modern Excel .xlsx as tab-separated rows per sheet."""
    try:
        import openpyxl
    except ImportError as e:
        raise RuntimeError("openpyxl is required to read .xlsx files") from e
    wb = openpyxl.load_workbook(full, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                out.append("\t".join("" if c is None else str(c) for c in row))
    wb.close()
    return "\n".join(out)


def _extract_xls(full: str) -> str:
    """Extract text from a LEGACY binary Excel .xls. Fix #6.

    Prefers xlrd (the only reader for the old BIFF format); falls back to pandas.
    """
    try:
        import xlrd
        book = xlrd.open_workbook(full)
        out = []
        for sh in book.sheets():
            out.append(f"# Sheet: {sh.name}")
            for rx in range(sh.nrows):
                vals = sh.row_values(rx)
                if any(v not in ("", None) for v in vals):
                    out.append("\t".join("" if v is None else str(v) for v in vals))
        return "\n".join(out)
    except ImportError:
        pass
    try:
        import pandas as pd
        sheets = pd.read_excel(full, sheet_name=None, engine="xlrd")
        out = []
        for name, df in sheets.items():
            out.append(f"# Sheet: {name}")
            out.append(df.to_csv(sep="\t", index=False))
        return "\n".join(out)
    except Exception as e:
        raise RuntimeError(
            "Reading legacy .xls requires the 'xlrd' package (pip install xlrd)."
        ) from e


def _extract_html(full: str) -> str:
    """Extract visible text from an HTML file. Fix #7 (127 HTML files)."""
    raw = open(full, encoding="utf-8", errors="replace").read()
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        import re as _re
        raw = _re.sub(r"(?is)<(script|style).*?</\1>", "", raw)
        return _re.sub(r"(?s)<[^>]+>", " ", raw)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _safe(path: str) -> str:
    full = os.path.normpath(os.path.join(_DATA_DIR, path))
    if not full.startswith(os.path.normpath(_DATA_DIR)):
        raise ValueError("Path escapes configured data_dir")
    if not os.path.exists(full):
        raise FileNotFoundError(f"Not found under data_dir: {path}")
    return full


# ── Tool functions ───────────────────────────────────────────────────────────

def list_documents(subfolder: str = "") -> List[dict]:
    """List readable documents under the configured document root, optionally
    restricted to a subfolder. Supports pdf/md/txt/csv/eml/json/html/docx/xls/xlsx."""
    root = _safe(subfolder) if subfolder else _DATA_DIR
    out: List[dict] = []
    for r, _, files in os.walk(root):
        for f in files:
            if f.lower().endswith(_SUPPORTED):
                p = os.path.join(r, f)
                out.append({
                    "path":  os.path.relpath(p, _DATA_DIR),
                    "bytes": os.path.getsize(p),
                })
    return out


def extract_text(path: str, max_chars: int = 20000) -> dict:
    """Extract plain text from a document given its path relative to the document
    root. Supports PDF, Word (.docx), Excel (.xls/.xlsx), HTML, and plain text
    formats (md/txt/csv/eml/json/log). Fix #6/#15."""
    full = _safe(path)
    low = full.lower()
    meta = {"pages": None, "pages_extracted": None}
    # G16/G17: prefer the shared, high-quality parser (core.document_parser) for
    # formats where it is materially better than the naive local path — PDF
    # (markitdown + hybrid OCR: preserves tables/amounts for scanned pages), and .doc /
    # .pptx / .rtf which the local reader can't handle at all. This unifies on ONE
    # good parser so the agent's reads match the KB's. Falls back to the local
    # extractor if the shared parser is unavailable.
    if low.endswith((".pdf",) + _CORE_EXT):
        try:
            from core.document_parser import parse_file
            ext = low.rsplit(".", 1)[-1]
            text = parse_file(full, ext, os.path.basename(full))
        except Exception as exc:
            # Fall back to the naive path only for PDF (the others have no local reader).
            if low.endswith(".pdf") and PdfReader is not None:
                reader = PdfReader(full)
                pages = min(len(reader.pages), _MAX_PAGES)
                text = "\n".join((reader.pages[i].extract_text() or "") for i in range(pages))
                meta = {"pages": len(reader.pages), "pages_extracted": pages}
            else:
                raise RuntimeError(f"Could not extract text from {os.path.basename(full)}: {exc}")
    elif low.endswith(_DOCX_EXT):
        text = _extract_docx(full)
    elif low.endswith(_XLSX_EXT):
        text = _extract_xlsx(full)
    elif low.endswith(_XLS_EXT):
        text = _extract_xls(full)
    elif low.endswith(_HTML_EXT):
        text = _extract_html(full)
    else:
        text = open(full, encoding="utf-8", errors="replace").read()
    return {
        "path":      path,
        "text":      text[:max_chars],
        "truncated": len(text) > max_chars,
        **meta,
    }


_MAX_BATCH_FILES = int(os.getenv("DOCUMENT_TOOLS_MAX_BATCH_FILES", "200"))


def extract_text_batch(paths: List[str], max_chars_each: int = 4000,
                       total_char_budget: int = 120000) -> dict:
    """Extract text from MANY documents in ONE call (Fix #7).

    Reading 100+ files one-at-a-time exhausts the agent's turn/context budget. This
    reads a list of files in a single call, capping each file's text (max_chars_each)
    and the overall payload (total_char_budget) so a large batch (e.g. 127 HTML files)
    can be summarised without ~254 round-trips. Files past the budget are reported as
    skipped so the agent can fetch them in a follow-up batch.

    Returns {documents: [{path, text, truncated, error?}], processed, skipped,
    total_requested}.
    """
    docs: List[dict] = []
    used = 0
    skipped: List[str] = []
    # G22: hard file-count cap so a runaway list (e.g. 5000 paths) can't tie up the
    # worker parsing every file. Excess files are reported as skipped for a follow-up.
    _paths = list(paths or [])
    if len(_paths) > _MAX_BATCH_FILES:
        skipped.extend(_paths[_MAX_BATCH_FILES:])
        _paths = _paths[:_MAX_BATCH_FILES]
    for p in _paths:
        if used >= total_char_budget:
            skipped.append(p)
            continue
        remaining = min(max_chars_each, total_char_budget - used)
        try:
            res = extract_text(p, max_chars=remaining)
            docs.append({"path": p, "text": res["text"], "truncated": res["truncated"]})
            used += len(res["text"])
        except Exception as e:  # noqa: BLE001 — surface per-file errors, keep going
            docs.append({"path": p, "text": "", "error": str(e)})
    return {
        "documents":       docs,
        "processed":       len(docs),
        "skipped":         skipped,
        "total_requested": len(paths),
        "note": (f"{len(skipped)} file(s) skipped — batch limit ({_MAX_BATCH_FILES} files) "
                 "or char budget reached. Call again with the skipped paths to continue.")
                 if skipped else "",
    }


def search_in_document(path: str, query: str, context_chars: int = 300) -> List[dict]:
    """Find occurrences of a query string (case-insensitive) inside a
    document; returns surrounding context snippets (max 20 hits)."""
    text = extract_text(path, max_chars=1_000_000)["text"]
    low, q, hits, start = text.lower(), query.lower(), [], 0
    while len(hits) < 20:
        i = low.find(q, start)
        if i < 0:
            break
        hits.append({
            "offset":  i,
            "snippet": text[max(0, i - context_chars // 2): i + context_chars // 2],
        })
        start = i + len(q)
    return hits
