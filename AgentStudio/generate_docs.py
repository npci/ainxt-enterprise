# SPDX-License-Identifier: MIT
"""
Generates a structured Word document describing the HITL, Loops,
and Existing Asset features of the ABStudio platform — strictly
based on the current implementation (no design changes).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---------- Styles ----------
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_kv_table(rows, headers=("Field", "Description")):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_heading('ABStudio Platform', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Feature Implementation Reference\n'
                'HITL  •  Loops  •  Existing Asset')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x33, 0x55, 0x88)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run('Version: Current Implementation Setup\nDate: 2026-06-15')
mr.italic = True
mr.font.size = Pt(11)

doc.add_paragraph()
add_para(
    'This document is a structured reference for three workflow node features '
    'in the ABStudio platform: Human-in-the-Loop (HITL), Loops, and Existing '
    'Asset (Sub-flow). For each feature, it explains what it does, which code '
    'files implement it, how data flows through the system, and a concrete '
    'walk-through example. Nothing in this document proposes new design — it '
    'mirrors the current code base only.'
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading('Table of Contents', level=1)
for i, item in enumerate([
    '1.  Platform Overview',
    '2.  Human-in-the-Loop (HITL)',
    '       2.1  What HITL Does',
    '       2.2  Implementation Files',
    '       2.3  HITL Modes',
    '       2.4  Data Flow',
    '       2.5  Worked Example',
    '3.  Loops',
    '       3.1  What the Loop Node Does',
    '       3.2  Implementation Files',
    '       3.3  Loop Modes (for_each / while / count)',
    '       3.4  Confidence Score & Continuation Contract',
    '       3.5  SSE Events Emitted',
    '       3.6  Worked Example — count mode with 3 iterations',
    '       3.7  Worked Example — while mode with confidence score',
    '       3.8  How to Increase the Confidence Score',
    '4.  Existing Asset (Sub-flow)',
    '       4.1  What the Existing Asset Node Does',
    '       4.2  Implementation Files',
    '       4.3  Node Configuration',
    '       4.4  Data Flow (agent vs workflow)',
    '       4.5  Worked Example',
    '5.  Feature Cross-reference',
]):
    add_para(item)

doc.add_page_break()

# ============================================================
# 1. PLATFORM OVERVIEW
# ============================================================
add_heading('1. Platform Overview', level=1)

add_para(
    'ABStudio executes user-built workflows on a native orchestration engine. '
    'A workflow is a directed graph of nodes (Agent, Condition, Loop, '
    'Subflow / Existing Asset, End). The engine streams progress to the '
    'frontend as Server-Sent Events (SSE).'
)
add_para('Key directories:', bold=True)
add_kv_table([
    ('backend/app/engine/native_engine.py',
     'Core graph walker; runs agents, loops, conditions, subflows, HITL.'),
    ('backend/app/engine/interface.py',
     'Defines SSE event names and loop-mode constants.'),
    ('backend/app/services/services.py',
     'Pure helpers — HITL mode parsing, prompt construction, condition eval.'),
    ('backend/app/tools/ask_human.py',
     'Declarative ask_human tool surface used by HITL.'),
    ('backend/app/checkpoint/store.py',
     'Persists HITL pause snapshots and per-node outputs.'),
    ('frontend/src/store/workflowStore.js',
     'Builds the execution payload (node defaults: maxIterations=3, etc.).'),
    ('frontend/src/features/workflows/editor/ConfigPanel.jsx',
     'UI editor for node settings (HITL mode, loop config, asset refId).'),
])

doc.add_page_break()

# ============================================================
# 2. HITL
# ============================================================
add_heading('2. Human-in-the-Loop (HITL)', level=1)

add_heading('2.1  What HITL Does', level=2)
add_para(
    'HITL pauses an agent\'s run so a human reviewer can approve, edit, or '
    'reject the next step before the workflow continues. It is configured '
    'per Agent node via the hitlMode field.'
)

add_heading('2.2  Implementation Files', level=2)
add_kv_table([
    ('backend/app/tools/ask_human.py',
     'Declares the ask_human tool shape. The engine intercepts ask_human calls '
     'by name and pauses — the tool itself is never actually executed.'),
    ('backend/app/services/services.py  (get_hitl_mode)',
     'Normalises legacy boolean / string hitlMode values into: "off", '
     '"after_response", "before_tool", "both".'),
    ('backend/app/engine/native_engine.py',
     'Implements the pause/snapshot/resume machinery. Key spots: '
     'HITL_SNAPSHOT_VERSION, the ask_human interception block (~line 1620), '
     'before_tool gate (~line 1661), and the resume() method.'),
    ('backend/app/checkpoint/store.py',
     'In-memory pause snapshot keyed by thread_id, plus the persisted '
     'Postgres variant in postgres_store.py.'),
    ('backend/app/api/chat.py',
     'SSE chat endpoint that delivers hitl_interrupt events to the client.'),
    ('frontend/src/features/workflows/editor/ConfigPanel.jsx',
     'HITL mode dropdown (off / before_tool / after_response / both).'),
])

add_heading('2.3  HITL Modes', level=2)
add_kv_table([
    ('off', 'No pausing. Agent runs to completion.'),
    ('after_response',
     'Agent finishes its reply, then pauses for human approval before the '
     'workflow advances to the next node.'),
    ('before_tool',
     'Agent decides on a tool call but the engine pauses BEFORE executing it. '
     'Human approves or rejects each tool call.'),
    ('both',
     'Pauses both before tool execution AND after the final response.'),
], headers=("Mode", "Behaviour"))

add_para(
    'Additionally, when hitlMode != "off", the engine registers the '
    'ask_human tool in that agent\'s function spec. The LLM can call '
    'ask_human voluntarily to ask the human a specific question with '
    '2–5 short options.', italic=True
)

add_heading('2.4  Data Flow', level=2)
add_para('1. The engine starts the agent and includes ask_human in its tools '
         'if HITL is enabled.', )
add_para('2. The agent either:')
add_bullet('emits an ask_human(question, options[], context) call, OR')
add_bullet('completes a final response (after_response mode), OR')
add_bullet('issues a regular tool call (before_tool mode).')
add_para('3. The engine intercepts, builds a snapshot containing:')
add_code('{\n'
         '  "version": 1,\n'
         '  "thread_id": "...",\n'
         '  "node_id": "agent_xxx",\n'
         '  "reason": "ask_human" | "before_tool" | "after_response",\n'
         '  "hitl_mode": "...",\n'
         '  "extra": {\n'
         '     "ask_human": {"question": "...", "options": [...], "context": "..."},\n'
         '     "tool_call_id": "..."\n'
         '  },\n'
         '  "state": <serialized execution state>\n'
         '}')
add_para('4. Saves the snapshot to the checkpoint store and emits an SSE '
         'hitl_interrupt event.')
add_para('5. The frontend renders the HITL card. The human selects an option / '
         'types a reply / approves / rejects.')
add_para('6. The /chat endpoint calls engine.resume(human_input). The engine '
         'reloads the snapshot, synthesises a tool result for the pending '
         'ask_human call (or a permit/deny for before_tool), and continues '
         'the ReAct loop from the same node.')

add_heading('2.5  Worked Example', level=2)
add_para('Setup: Agent "TravelBooker" with hitlMode="after_response".', bold=True)
add_para('User input:', bold=True)
add_code('"Book me a flight to Tokyo next Friday under $1200."')
add_para('Agent reply (the LLM finishes):', bold=True)
add_code('"I found ANA flight NH7 departing 09:50, return Sun 17:40, total $1,148. '
         'Shall I confirm?"')
add_para('Engine behaviour:', bold=True)
add_bullet('Detects hitlMode="after_response".')
add_bullet('Persists snapshot {reason: "after_response", node_id: travel_booker}.')
add_bullet('Emits hitl_interrupt → frontend renders approve / reject card.')
add_para('Human clicks "Approve":', bold=True)
add_bullet('POST /chat-resume with body "Approve".')
add_bullet('engine.resume() loads snapshot, marks the agent\'s output as accepted, '
           'unwinds the pause flag, and advances to the next node.')
add_para('Human clicks "Reject" with note "Try Singapore Airlines":', bold=True)
add_bullet('Engine re-enters the agent loop on the same node with the '
           'rejection note appended to the conversation.')
add_bullet('Agent produces a new draft using SQ instead of NH.')

doc.add_page_break()

# ============================================================
# 3. LOOPS
# ============================================================
add_heading('3. Loops', level=1)

add_heading('3.1  What the Loop Node Does', level=2)
add_para(
    'A Loop node wraps a subgraph (the "body") and re-runs it until a '
    'termination condition is met. The loop node has two output handles: '
    '"body" (where iteration goes) and "exit" (where control flows once '
    'the loop ends).'
)

add_heading('3.2  Implementation Files', level=2)
add_kv_table([
    ('backend/app/engine/native_engine.py  (_run_loop, ~line 2370)',
     'The main loop driver. Reads node config, walks the body subgraph using '
     '_traverse(stop_at={node_id}), emits all loop_* SSE events.'),
    ('backend/app/engine/native_engine.py  (_build_loop_directive, ~line 2590)',
     'Appends "Loop context" and (for while-mode) the "Loop continuation '
     'contract" block to a body-agent\'s prompt so it knows it must emit '
     'a score and a changes string.'),
    ('backend/app/engine/interface.py',
     'Defines LOOP_MODE_FOR_EACH / LOOP_MODE_WHILE / LOOP_MODE_COUNT and the '
     'loop SSE events: loop_iteration_start, loop_iteration_end, '
     'loop_condition_eval, loop_iteration_summary, loop_final_summary, '
     'loop_complete.'),
    ('backend/app/services/services.py',
     'build_expression_from_case + evaluate_condition — shared condition DSL '
     'used by while-mode loops.'),
    ('backend/app/checkpoint/store.py',
     'Persists each node\'s last output keyed by (thread_id, node_id) so the '
     'Loop config picker can show real upstream lists.'),
    ('frontend/src/store/workflowStore.js',
     'Loop node defaults: mode=for_each, itemsExpression="input.items", '
     'count=3, maxIterations=3, iteratorVar="item".'),
    ('frontend/src/features/workflows/editor/ConfigPanel.jsx + LoopItemsPicker.jsx',
     'UI editor for loop mode, count, items expression, cases, and the '
     'connection-aware list picker.'),
])

add_heading('3.3  Loop Modes (for_each / while / count)', level=2)
add_kv_table([
    ('for_each',
     'Iterates over a list resolved from itemsExpression (default '
     '"input.items"). Each iteration sets loop.item to the current element. '
     'Stops when the list is exhausted.'),
    ('count',
     'Iterates a fixed number of times. count is read from the node config '
     '(default 3). loop.item is the integer index.'),
    ('while',
     'Re-iterates while one of the user-configured cases evaluates true. '
     'Cases use the same condition DSL as the Condition node and read from '
     'the agent\'s output (e.g. input.score < 0.7). Stops when no case '
     'matches OR when maxIterations is hit.'),
], headers=("Mode", "Behaviour"))

add_para('Key config fields on a loop node:', bold=True)
add_kv_table([
    ('mode',              'for_each | while | count (default: for_each).'),
    ('count',             'Fixed iteration count for count-mode (default: 3).'),
    ('itemsExpression',   'Dotted path to a list for for_each (default: "input.items").'),
    ('iteratorVar',       'Variable name exposed inside body prompts (default: "item").'),
    ('cases',             'Continuation expressions for while-mode.'),
    ('maxIterations',     'Hard safety ceiling for ALL modes (default: 3 in the UI; '
                          'falls back to 25 inside the engine if not set).'),
])

add_heading('3.4  Confidence Score & Continuation Contract', level=2)
add_para(
    'When the loop is in while-mode and at least one case references a field, '
    '_build_loop_directive injects a "Loop continuation contract" into the '
    'body agent\'s prompt. This is how the system gets a "confidence score" '
    'out of the agent on every iteration.'
)

add_para('What the agent is forced to emit on the LAST line of its response:', bold=True)
add_code('{"score": 0.0..1.0, "changes": "<one-line summary of what you '
         'changed this round>"}')

add_para('Where the score comes from:', bold=True)
add_bullet('The score is the agent\'s OWN self-rating of the quality / '
           'confidence of its current artifact.')
add_bullet('The engine does not compute the score — it scrapes it from the '
           'agent\'s final JSON line via resolve_routing_state().')
add_bullet('Once scraped, the loop\'s while-mode expression (e.g. '
           '"input.score > 0.7") is evaluated. If it still matches, the body '
           'runs again. If it no longer matches, the loop exits through the '
           '"exit" handle.')

add_para('Aggregation across iterations:', bold=True)
add_bullet('iter_summaries is a list of {index, score, changes} records — one '
           'per iteration where a numeric score was produced.')
add_bullet('initial_score = first iteration\'s score.')
add_bullet('final_score   = last iteration\'s score.')
add_bullet('delta         = final_score - initial_score (rounded to 4 dp).')
add_bullet('All of this is shipped to the UI in a single loop_final_summary '
           'SSE event, which renders the end-of-loop chat bubble.')

add_heading('3.5  SSE Events Emitted', level=2)
add_kv_table([
    ('loop_iteration_start',
     '{node_id, mode, index, total?} — emitted at the start of every iteration.'),
    ('loop_condition_eval',
     '{node_id, index, case_results, will_continue, eval_state?} — fires once '
     'per round in while-mode showing which case matched.'),
    ('loop_iteration_summary',
     '{node_id, index, score, changes, output_preview} — only emitted in '
     'while-mode (drives the per-row "Confidence Score" pill in the UI).'),
    ('loop_iteration_end',
     '{node_id, index} — fired at the end of each iteration.'),
    ('loop_final_summary',
     '{iterations[], initial_score, final_score, delta, final_output, '
     'final_structured?, max_iterations_hit} — the end-of-loop bubble.'),
    ('loop_complete',
     '{node_id, total_iterations, max_iterations_hit} — terminal event.'),
])

add_heading('3.6  Worked Example — count mode with 3 iterations', level=2)
add_para('Setup:', bold=True)
add_bullet('Loop node: mode="count", count=3, maxIterations=3.')
add_bullet('Body: a single Agent node "Refiner" that improves a poem draft.')
add_bullet('Initial input passed in: "Write a 4-line poem about the monsoon."')

add_para('Iteration log (what is emitted on the SSE stream):', bold=True)
add_code(
    'loop_iteration_start  {index: 0, total: 3, mode: "count"}\n'
    '  agent_start         Refiner\n'
    '  agent_complete      "Rain taps the tin, the city breathes ..." (draft 1)\n'
    'loop_iteration_end    {index: 0}\n'
    '\n'
    'loop_iteration_start  {index: 1, total: 3}\n'
    '  agent_start         Refiner\n'
    '  agent_complete      "Rain plays a slow drum on the tin roofs ..." (draft 2)\n'
    'loop_iteration_end    {index: 1}\n'
    '\n'
    'loop_iteration_start  {index: 2, total: 3}\n'
    '  agent_start         Refiner\n'
    '  agent_complete      "Drums on tin, a city exhales in green ..." (draft 3)\n'
    'loop_iteration_end    {index: 2}\n'
    '\n'
    'loop_final_summary    {iterations: [], initial_score: null, final_score: null,\n'
    '                       max_iterations_hit: false, final_output: "Drums on tin..."}\n'
    'loop_complete         {total_iterations: 3}\n'
)

add_para('What "data goes back":', bold=True)
add_bullet('After each iteration, state.current_input is overwritten by the '
           'Refiner\'s output. Iteration N reads iteration N-1\'s output as '
           'its prompt context (via build_agent_prompt + execution_trace).')
add_bullet('The body agent prompt also receives a "Loop context" trailer with '
           'loop.index, loop.total, loop.item, loop.var so it knows which '
           'round it is on.')
add_bullet('Because this is count-mode, NO continuation contract is appended, '
           'so the agent does NOT have to emit a JSON score line. '
           'iter_summaries stays empty and initial_score / final_score are null.')

add_heading('3.7  Worked Example — while mode with confidence score', level=2)
add_para('Setup:', bold=True)
add_bullet('Loop node: mode="while", maxIterations=3.')
add_bullet('Case configured: input.score < 0.8  (i.e. keep iterating while '
           'self-rated score is under 0.8).')
add_bullet('Body: Agent node "Polisher" that refines a summary.')

add_para('What the engine appends to Polisher\'s instructions:', bold=True)
add_code(
    '## Loop continuation contract (REQUIRED)\n'
    '\n'
    'This agent runs inside a loop that re-iterates while one of these\n'
    'expressions is true:\n'
    '  - input.score < 0.8\n'
    '\n'
    'On the LAST line of your response, emit a single JSON object with\n'
    'exactly these keys (no markdown fence, no commentary after it):\n'
    '  {"score": <number 0..1>, "changes": "<one-line summary of what\n'
    '  you changed this round>"}\n'
)

add_para('Round-by-round trace:', bold=True)
add_code(
    'iter 0  → Polisher returns draft + final line:\n'
    '           {"score": 0.55, "changes": "initial draft"}\n'
    '         loop_iteration_summary {index:0, score:0.55, changes:"initial draft"}\n'
    '         Continuation check: 0.55 < 0.8 → TRUE → iterate again.\n'
    '\n'
    'iter 1  → Polisher returns refined draft + final line:\n'
    '           {"score": 0.72, "changes": "tightened intro, fixed two facts"}\n'
    '         loop_iteration_summary {index:1, score:0.72, changes:"..."}\n'
    '         Continuation check: 0.72 < 0.8 → TRUE → iterate again.\n'
    '\n'
    'iter 2  → Polisher returns final draft + final line:\n'
    '           {"score": 0.85, "changes": "added executive summary"}\n'
    '         loop_iteration_summary {index:2, score:0.85, changes:"..."}\n'
    '         Continuation check: 0.85 < 0.8 → FALSE → exit loop.\n'
    '\n'
    'loop_final_summary {\n'
    '   iterations: [\n'
    '       {index:0, score:0.55, changes:"initial draft"},\n'
    '       {index:1, score:0.72, changes:"tightened intro..."},\n'
    '       {index:2, score:0.85, changes:"added executive summary"},\n'
    '   ],\n'
    '   initial_score: 0.55,\n'
    '   final_score:   0.85,\n'
    '   delta:         0.30,\n'
    '   max_iterations_hit: false\n'
    '}\n'
    'loop_complete {total_iterations: 3, max_iterations_hit: false}\n'
)

add_para('What happens if maxIterations=3 but score never crosses 0.8:', bold=True)
add_bullet('iteration index reaches 3 → hit_safety_cap = True → break.')
add_bullet('loop_final_summary fires with max_iterations_hit: true.')
add_bullet('The UI shows a banner explaining the loop hit the safety cap '
           'before the natural stop condition.')

add_heading('3.8  How to Increase the Confidence Score', level=2)
add_para(
    'The "confidence score" is the score field the body agent self-reports '
    'on its last line. There is no separate scorer agent; the loop trusts '
    'the body agent\'s number. To make scores rise across iterations:'
)
add_bullet('Phrase the case so it RAISES the bar — e.g. input.score < 0.8 '
           'forces the agent to keep going until it self-rates ≥ 0.8.')
add_bullet('Give the body agent stronger instructions: tell it exactly what '
           'a "high-confidence" answer looks like for your domain.')
add_bullet('Increase maxIterations so a slow climber has room to converge '
           '(but remember it is a hard ceiling).')
add_bullet('Pair the loop with tools (KB search, validators) so each '
           'iteration can ground its score in fresh evidence rather than '
           'just paraphrasing the previous draft.')
add_bullet('Inspect loop_iteration_summary.changes — if "changes" is empty '
           'or repetitive, the agent isn\'t actually improving, and the '
           'instructions need to be tightened.')

doc.add_page_break()

# ============================================================
# 4. EXISTING ASSET
# ============================================================
add_heading('4. Existing Asset (Sub-flow)', level=1)

add_heading('4.1  What the Existing Asset Node Does', level=2)
add_para(
    'The Existing Asset node lets a workflow embed and reuse a previously-'
    'saved Agent or another Workflow as a single graph node. Internally it '
    'is a "subflow" node (UI label: "Existing Asset"). When the engine hits '
    'it, _run_subflow dispatches into the saved asset and forwards its '
    'output back into the parent workflow.'
)

add_heading('4.2  Implementation Files', level=2)
add_kv_table([
    ('backend/app/engine/native_engine.py  (_run_subflow, ~line 1847)',
     'Looks up the linked asset by refId, runs it (AgentRunner.run for kind='
     '"agent", recursive engine.execute for kind="workflow"), forwards SSE '
     'events with the parent label prefixed, and pipes the output back into '
     'state.current_input.'),
    ('backend/app/engine/native_engine.py  (subflow_stack)',
     'Per-run recursion guard. Detects "Sub-flow loop detected" when an '
     'asset transitively references itself.'),
    ('backend/agent_factory/pipeline.py  (AgentRunner)',
     'Same runner used by /agent-runner/chat. Reuses per-user credentials, '
     'tool dispatch, and persistence.'),
    ('frontend/src/features/workflows/editor/nodes/SubflowNode.jsx',
     'Stacked-squares icon and inline asset picker on the canvas.'),
    ('frontend/src/features/workflows/editor/Sidebar.jsx',
     'Drag source labelled "Existing Asset".'),
    ('frontend/src/store/workflowStore.js',
     'Exports {id, type:"subflow", kind, refId, refName} for the subflow node; '
     'guards against unlinked nodes ("An \'Existing Asset\' node is not '
     'linked to a workflow or agent yet").'),
])

add_heading('4.3  Node Configuration', level=2)
add_kv_table([
    ('kind',    '"agent" or "workflow" — which catalog to look the asset up in.'),
    ('refId',   'The saved asset\'s id. Engine refuses to run if this is empty.'),
    ('refName', 'Cached human-readable label, used for SSE attribution and '
                'breadcrumbs like "ParentWF ▸ ChildAgent".'),
])

add_heading('4.4  Data Flow (agent vs workflow)', level=2)
add_para('kind = "agent":', bold=True)
add_bullet('Engine instantiates AgentRunner(AgentRegistry, MonitoringLogger).')
add_bullet('Calls runner.run(refId, state.current_input, user_id, email, '
           'department, is_admin).')
add_bullet('Emits agent_start (if this subflow is the workflow\'s final node) '
           'or agent_progress otherwise.')
add_bullet('Adds {agent: refName, output, node_id} to state.execution_trace '
           'and pushes any generated files into state.generated_files.')
add_bullet('Updates state.current_input = response so the next node in the '
           'parent graph receives the sub-agent\'s reply.')

add_para('kind = "workflow":', bold=True)
add_bullet('Loads the referenced workflow\'s graphData and recursively calls '
           'self.execute() on it.')
add_bullet('Wraps inner SSE agent names as "<parentRefName> ▸ <innerAgent>" '
           'so the live trace shows the call hierarchy.')
add_bullet('If the inner run pauses for HITL, the engine persists a '
           'snapshot with reason="subflow_pending" and the parent suspends '
           'until /chat-resume is called.')

add_para('Recursion guard:', bold=True)
add_bullet('subflow_stack is a per-run list of "{kind}:{refId}" guard keys.')
add_bullet('If the same key appears twice in the stack the engine emits an '
           'error event ("Sub-flow loop detected") instead of recursing.')

add_heading('4.5  Worked Example', level=2)
add_para('Setup:', bold=True)
add_bullet('Parent workflow: "QuarterlyReport".')
add_bullet('Three nodes wired sequentially: '
           'DataFetcher (Agent) → ExistingAsset (kind="workflow", '
           'refName="RevenueAnalyzer") → Summariser (Agent).')

add_para('Walk-through:', bold=True)
add_code(
    '1. DataFetcher runs, fetches Q1 figures.\n'
    '   state.current_input = "Q1 revenue rows: ..."\n'
    '\n'
    '2. Engine reaches the ExistingAsset node.\n'
    '   _run_subflow detects kind="workflow", refId="wf_rev_an".\n'
    '   Pushes "workflow:wf_rev_an" onto subflow_stack.\n'
    '   Loads the RevenueAnalyzer graphData and recursively executes it.\n'
    '   Inner SSE events are forwarded as:\n'
    '       agent_start  {agent: "RevenueAnalyzer ▸ Cleaner"}\n'
    '       agent_start  {agent: "RevenueAnalyzer ▸ Scorer"}\n'
    '       agent_complete {agent: "RevenueAnalyzer ▸ Scorer", output: "...JSON..."}\n'
    '   Inner workflow returns final output → parent\'s state.current_input is\n'
    '   updated with that output. subflow_stack pops "workflow:wf_rev_an".\n'
    '\n'
    '3. Summariser runs with the RevenueAnalyzer output as its input and\n'
    '   produces the final exec summary that ships to the user.\n'
)

add_para('If RevenueAnalyzer paused for HITL inside its own graph:', bold=True)
add_bullet('Parent run is suspended with reason="subflow_pending".')
add_bullet('snapshot.extra carries inner_thread_id, subflow_ref_id, '
           'subflow_ref_name, subflow_kind.')
add_bullet('When the human resumes, the engine restores the inner thread, '
           'finishes the inner workflow, and then continues the parent '
           'traversal from the subflow node\'s successor (Summariser).')

doc.add_page_break()

# ============================================================
# 5. CROSS-REFERENCE
# ============================================================
add_heading('5. Feature Cross-reference', level=1)

add_kv_table([
    ('HITL',
     'Per-Agent feature (hitlMode field).',
     'ask_human.py + native_engine.py (intercept block, resume()) + checkpoint store.',
     'hitl_interrupt, hitl_resumed'),
    ('Loop',
     'Standalone node type "loop" with body / exit handles.',
     'native_engine._run_loop, _build_loop_directive, interface.LOOP_MODES.',
     'loop_iteration_start, loop_condition_eval, loop_iteration_summary, '
     'loop_iteration_end, loop_final_summary, loop_complete'),
    ('Existing Asset',
     'Standalone node type "subflow" — drag from sidebar as "Existing Asset".',
     'native_engine._run_subflow + AgentRunner + recursive engine.execute.',
     'Forwards inner agent_start / agent_complete with prefixed names; may '
     'pause with reason=subflow_pending'),
], headers=("Feature", "Surface", "Core code", "Key SSE events"))

doc.add_paragraph()
add_para('End of document.', italic=True)

# ---------- Save ----------
out_path = "D:/ainxt-platform/ABStudio/ABStudio_Features_HITL_Loops_ExistingAsset.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
