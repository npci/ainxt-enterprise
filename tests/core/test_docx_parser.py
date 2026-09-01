# SPDX-License-Identifier: Apache-2.0
# ============================================================
# tests/core/test_docx_parser.py
#
# Unit tests for core.document_parser.parse_docx()
# — the server-side DOCX parser used by ABStudio's Agent Runner
#   upload path (documents.py → ocr_pipeline.py) and the desktop
#   Buddy extract_document tool.
#
# These tests build real .docx files in a temp directory using
# python-docx, then (for the regression cases) surgically wrap
# the body's <w:p>/<w:tbl> elements in <w:sdt><w:sdtContent>...
# to reproduce Word's "Content Control" feature — used heavily by
# templated policy/finance manuals, government forms, and DMS-
# managed documents for version blocks, approval fields, and
# placeholder text.
#
# Regression covered:
#   Before the fix, parse_docx() only walked doc.paragraphs /
#   doc.tables, which python-docx populates from TOP-LEVEL
#   <w:p>/<w:tbl> elements only. Anything nested inside a
#   <w:sdt> (Structured Document Tag) wrapper was invisible to
#   both doc.paragraphs and doc.tables, so a document built
#   entirely — or mostly — out of content controls silently
#   parsed to "" with no exception and no sentinel. Upstream,
#   that surfaced as the misleading error "File parsed but
#   contained no readable text."
#
# Coverage:
#   1. Normal document (no sdt)      → baseline, unaffected by fix
#   2. Single sdt-wrapped paragraph  → must now be extracted
#   3. Nested sdt (sdt-in-sdt)       → recursion must descend fully
#   4. sdt-wrapped table             → table content must be found
#   5. Mixed: sdt + top-level content, ordering preserved
#   6. Malformed/deeply-nested sdt   → depth guard must not crash
# ============================================================

from __future__ import annotations

from pathlib import Path

import pytest

# Skip the entire module if python-docx is not installed
# (it is in requirements.txt but may be absent in a minimal CI env).
docx = pytest.importorskip("docx")

from core.document_parser import parse_docx  # noqa: E402


# ── helpers ──────────────────────────────────────────────────────────────────

def _wrap_in_sdt(element) -> None:
    """
    Surgically rewrite a top-level body element (a <w:p> or <w:tbl>) in
    place so it is wrapped as:

        <w:sdt><w:sdtPr/><w:sdtContent>{element}</w:sdtContent></w:sdt>

    This mimics what Word does when a user inserts a Content Control
    (Developer tab → Rich Text / Plain Text / Table content control)
    around existing body content.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    body = element.getparent()
    idx = list(body).index(element)

    nsmap = element.nsmap
    sdt = etree.SubElement(body, qn("w:sdt"))
    etree.SubElement(sdt, qn("w:sdtPr"))
    sdt_content = etree.SubElement(sdt, qn("w:sdtContent"))

    body.remove(element)
    sdt_content.append(element)

    body.remove(sdt)
    body.insert(idx, sdt)


def _wrap_in_nested_sdt(element, depth: int = 2) -> None:
    """Wrap `element` in `depth` levels of nested <w:sdt><w:sdtContent>."""
    for _ in range(depth):
        _wrap_in_sdt(element)
        # The element we just wrapped is now inside a fresh sdtContent;
        # to nest again we need to wrap the *sdt* we just created, not
        # the original element, so find it again.
        element = element.getparent()  # sdtContent
        element = element.getparent()  # sdt


def _make_docx_with_paragraphs(path: Path, paragraphs: list[str]) -> "docx.Document":
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return doc


# ── fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture()
def tmp(tmp_path):
    return tmp_path


# ── tests ─────────────────────────────────────────────────────────────────────

class TestParseDocxBaseline:
    """Normal documents without content controls — must be unaffected by the fix."""

    def test_simple_paragraphs(self, tmp):
        p = tmp / "simple.docx"
        _make_docx_with_paragraphs(p, ["Hello world", "Second paragraph"])
        result = parse_docx(str(p))
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_heading_and_table(self, tmp):
        from docx import Document

        p = tmp / "heading_table.docx"
        doc = Document()
        doc.add_heading("Report Title", level=1)
        doc.add_paragraph("Intro text.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Score"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "95"
        doc.save(str(p))

        result = parse_docx(str(p))
        assert "# Report Title" in result
        assert "Intro text." in result
        assert "Name" in result and "Score" in result
        assert "Alice" in result and "95" in result


class TestParseDocxContentControls:
    """
    Regression tests for the <w:sdt> fix. Before the fix, all of these
    would return "" (or an empty-looking result) because doc.paragraphs
    and doc.tables cannot see sdt-nested content.
    """

    def test_single_sdt_wrapped_paragraph(self, tmp):
        p = tmp / "sdt_single.docx"
        doc = _make_docx_with_paragraphs(p, ["This is inside a content control."])
        _wrap_in_sdt(doc.paragraphs[0]._p)
        doc.save(str(p))

        result = parse_docx(str(p))
        assert "This is inside a content control." in result

    def test_document_entirely_content_controls(self, tmp):
        """A doc built entirely out of sdt-wrapped paragraphs must not be empty."""
        p = tmp / "sdt_all.docx"
        doc = _make_docx_with_paragraphs(
            p, ["Version: 1.0", "Approved by: J. Doe", "Effective date: 2026-01-01"]
        )
        for para in list(doc.paragraphs):
            _wrap_in_sdt(para._p)
        doc.save(str(p))

        result = parse_docx(str(p))
        assert result.strip() != ""
        assert "Version: 1.0" in result
        assert "Approved by: J. Doe" in result
        assert "Effective date: 2026-01-01" in result

    def test_nested_sdt(self, tmp):
        """sdt-in-sdt (e.g. a content control inside a repeating section) must recurse fully."""
        p = tmp / "sdt_nested.docx"
        doc = _make_docx_with_paragraphs(p, ["Deeply nested text."])
        _wrap_in_nested_sdt(doc.paragraphs[0]._p, depth=3)
        doc.save(str(p))

        result = parse_docx(str(p))
        assert "Deeply nested text." in result

    def test_sdt_wrapped_table(self, tmp):
        from docx import Document

        p = tmp / "sdt_table.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Policy No"
        table.cell(1, 1).text = "FIN-2026-27"
        _wrap_in_sdt(table._tbl)
        doc.save(str(p))

        result = parse_docx(str(p))
        assert "Field" in result and "Value" in result
        assert "Policy No" in result and "FIN-2026-27" in result

    def test_mixed_top_level_and_sdt_preserves_order(self, tmp):
        p = tmp / "sdt_mixed.docx"
        doc = _make_docx_with_paragraphs(
            p, ["First (top-level).", "Second (will be wrapped).", "Third (top-level)."]
        )
        _wrap_in_sdt(doc.paragraphs[1]._p)
        doc.save(str(p))

        result = parse_docx(str(p))
        first_idx = result.find("First (top-level).")
        second_idx = result.find("Second (will be wrapped).")
        third_idx = result.find("Third (top-level).")
        assert first_idx != -1 and second_idx != -1 and third_idx != -1
        assert first_idx < second_idx < third_idx

    def test_deep_nesting_does_not_crash(self, tmp):
        """Pathological/malformed nesting must hit the depth guard, not crash or hang."""
        p = tmp / "sdt_deep.docx"
        doc = _make_docx_with_paragraphs(p, ["Text past the depth guard."])
        _wrap_in_nested_sdt(doc.paragraphs[0]._p, depth=30)
        doc.save(str(p))

        # Must not raise; content beyond _MAX_SDT_DEPTH is simply not reached.
        result = parse_docx(str(p))
        assert isinstance(result, str)


class TestParseDocxOutputFormat:
    """Output format matches what upstream (ocr_pipeline / gateway) expects."""

    def test_output_is_string(self, tmp):
        p = tmp / "fmt.docx"
        _make_docx_with_paragraphs(p, ["Some text."])
        result = parse_docx(str(p))
        assert isinstance(result, str)

    def test_no_binary_garbage(self, tmp):
        p = tmp / "clean.docx"
        _make_docx_with_paragraphs(p, ["AiNxt clean text"])
        result = parse_docx(str(p))
        assert "\x00" not in result
        assert result.encode("utf-8").decode("utf-8") == result

    def test_missing_file_returns_error_string_not_exception(self, tmp):
        result = parse_docx(str(tmp / "does_not_exist.docx"))
        assert isinstance(result, str)
        assert "[DOCX parse error" in result
